# Export LOs and Questions to a .docx file
from io import BytesIO
from typing import List, Dict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

def build_docx(los:List[dict], questions_by_lo:Dict[str,list])->bytes:
    doc=Document()
    # Force Title and Heading styles to use Arial as well
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name].font
        style.name = "Arial"
        #style.size = Pt(11)

    # Normal font is Ariel 11pt
    font=doc.styles['Normal'].font
    font.name="Arial"
    font.size=Pt(11)

    # Title line
    doc.add_paragraph("Assessment Questions", style="Title")

    # Go over Learning Objectives
    for lo in los:
        lo_id=lo["id"]
        final=lo.get("final_text")
        level=lo["intended_level"]

        # LO heading
        doc.add_heading(f"Learning Objective: {final}", level=2)
        # Bloom line is italic
        para = doc.add_paragraph()
        run = para.add_run(f"Bloom level: {level}")
        run.italic = True

        # Go over questions
        qs=questions_by_lo.get(lo_id,[])
        for idx,q in enumerate(qs,1):
            # Question stem
            para = doc.add_paragraph()
            run = para.add_run(f"{idx}. {q['stem']}")
            run.bold = True
            run.font.color.rgb = RGBColor(0x5F, 0x49, 0x7A)  # purple-like color

            # Options with yellow highlight for the correct one
            correct=q.get("correct_option_id")
            for opt in q["options"]:
                para = doc.add_paragraph()
                run = para.add_run(f"   ({opt['id']}) {opt['text']}")
                if opt["id"]==correct:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            # Answer line
            para = doc.add_paragraph()
            run = para.add_run("Answer: ")
            run.bold = True
            para.add_run(correct)
           
            # Feedback block
            para = doc.add_paragraph()
            run = para.add_run(f"Feedback:")
            run.bold = True
            for opt in q["options"]:
                doc.add_paragraph(f"   ({opt['id']}) {opt['option_rationale']}")
            
            # Content reference
            para = doc.add_paragraph()
            run = para.add_run("Content reference: ")
            run.bold = True
            para.add_run(q['contentReference'])

            # Cognitive rationale
            para = doc.add_paragraph()
            run = para.add_run("Rationale for Bloom level: ")
            run.bold = True
            para.add_run(q['cognitive_rationale'])
    
    bio=BytesIO()
    doc.save(bio)
    return bio.getvalue()
