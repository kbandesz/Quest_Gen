# Export LOs and Questions to a .docx file
from io import BytesIO
from typing import List, Dict, Optional, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

def build_questions_docx(los:List[dict], questions_by_lo:Dict[str,list], include:Optional[Dict[str,bool]]=None)->bytes:
    doc = Document()
    # Force Title and Heading styles to use Arial as well
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name].font
        style.name = "Arial"

    # Normal font is Arial 11pt
    font = doc.styles['Normal'].font
    font.name = "Arial"
    font.size = Pt(11)

    # Title line
    doc.add_paragraph("Assessment Questions", style="Title")

    # Default include options (all True)
    if include is None:
        include = {
            "los": True,
            "bloom": True,
            "answer": True,
            "feedback": True,
            "content": True,
            "rationale": True,
        }

    # Go over Learning Objectives
    for lo in los:
        lo_id = lo.get("id")
        final = lo.get("final_text")
        level = lo.get("intended_level")

        # LO heading
        if include.get("los", True):
            doc.add_heading(f"Learning Objective: {final}", level=2)

        # Bloom line is italic
        if include.get("bloom", True):
            para = doc.add_paragraph()
            run = para.add_run(f"Bloom level: {level}")
            run.italic = True

        # Go over questions
        qs = questions_by_lo.get(lo_id, [])
        for idx, q in enumerate(qs, 1):
            # Question stem
            para = doc.add_paragraph()
            run = para.add_run(f"{idx}. {q.get('stem','')}")
            run.bold = True
            run.font.color.rgb = RGBColor(0x5F, 0x49, 0x7A)  # purple-like color

            # Options with yellow highlight for the correct one
            correct = q.get("correct_option_id")
            for opt in q.get("options", []):
                para = doc.add_paragraph()
                run = para.add_run(f"   ({opt.get('id')}) {opt.get('text','')}")
                if opt.get("id") == correct and include.get("answer", True):
                    try:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    except Exception:
                        # Some python-docx versions may not support highlight on run.font
                        pass

            # Answer line
            if include.get("answer", True):
                para = doc.add_paragraph()
                run = para.add_run("Answer: ")
                run.bold = True
                para.add_run(str(correct))

            # Feedback block
            if include.get("feedback", True):
                para = doc.add_paragraph()
                run = para.add_run("Feedback:")
                run.bold = True
                for opt in q.get("options", []):
                    doc.add_paragraph(f"   ({opt.get('id')}) {opt.get('option_rationale','')}")

            # Content reference
            if include.get("content", True):
                para = doc.add_paragraph()
                run = para.add_run("Content reference: ")
                run.bold = True
                para.add_run(q.get('contentReference',''))

            # Cognitive rationale
            if include.get("rationale", True):
                para = doc.add_paragraph()
                run = para.add_run("Rationale for Bloom level: ")
                run.bold = True
                para.add_run(q.get('cognitive_rationale',''))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_outline_docx(outline: Dict[str, Any]) -> bytes:
    """Construct a polished DOCX document summarizing the course outline."""

    doc = Document()

    # Harmonize font selections for a cohesive look
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = doc.styles[style_name].font
        style.name = "Calibri"

    normal_font = doc.styles["Normal"].font
    normal_font.name = "Calibri"
    normal_font.size = Pt(11)

    def _add_section_title(text: str, *, level: int = 2, accent: Optional[RGBColor] = None):
        heading = doc.add_heading(text, level=level)
        if accent:
            for run in heading.runs:
                run.font.color.rgb = accent

    course_title = outline.get("courseTitle") or "Untitled Course"

    doc.add_paragraph("Course Outline", style="Title")
    _add_section_title(course_title, level=1, accent=RGBColor(0x2F, 0x55, 0x7F))

    course_objectives = outline.get("courseLevelObjectives") or []
    if course_objectives:
        _add_section_title("Course-Level Objectives", level=2, accent=RGBColor(0x2F, 0x55, 0x7F))
        for obj in course_objectives:
            para = doc.add_paragraph(obj, style="List Bullet")
            if para.runs:
                para.runs[0].font.size = Pt(11)

    modules = outline.get("modules") or []
    for module_index, module in enumerate(modules, 1):
        module_title = module.get("moduleTitle") or f"Module {module_index}"
        overview = module.get("overview")

        _add_section_title(
            f"Module {module_index}: {module_title}",
            level=2,
            accent=RGBColor(0x5F, 0x49, 0x7A),
        )
        if overview:
            para = doc.add_paragraph(overview)
            para_format = para.paragraph_format
            para_format.space_after = Pt(6)

        sections = module.get("sections") or []
        for section_index, section in enumerate(sections, 1):
            section_title = section.get("sectionTitle") or f"Section {module_index}.{section_index}"
            _add_section_title(
                f"Section {module_index}.{section_index}: {section_title}",
                level=3,
                accent=RGBColor(0x00, 0x66, 0x66),
            )

            section_objectives = section.get("sectionLevelObjectives") or []
            if section_objectives:
                highlight = doc.add_paragraph("Section Objectives:")
                if highlight.runs:
                    highlight.runs[0].bold = True
                for obj in section_objectives:
                    doc.add_paragraph(obj, style="List Bullet")

            units = section.get("units") or []
            if units:
                table = doc.add_table(rows=1, cols=2)
                table.style = "Light Grid Accent 1"
                header_cells = table.rows[0].cells
                header_cells[0].text = "Unit"
                header_cells[1].text = "Key Points"
                for cell in header_cells:
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True

                for unit_index, unit in enumerate(units, 1):
                    unit_title = unit.get("unitTitle") or f"Unit {module_index}.{section_index}.{unit_index}"
                    unit_objective = unit.get("unitLevelObjective")
                    key_points = unit.get("keyPoints") or []

                    row_cells = table.add_row().cells
                    unit_cell = row_cells[0].paragraphs[0]
                    unit_run = unit_cell.add_run(f"{module_index}.{section_index}.{unit_index} {unit_title}")
                    unit_run.bold = True
                    if unit_objective:
                        objective_run = unit_cell.add_run("\nObjective: ")
                        objective_run.bold = True
                        unit_cell.add_run(unit_objective)

                    points_cell = row_cells[1].paragraphs[0]
                    if key_points:
                        for point in key_points:
                            run = points_cell.add_run(f"• {point}\n")
                            run.font.size = Pt(10)
                    else:
                        points_cell.add_run("No key points provided.")

                doc.add_paragraph("")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

